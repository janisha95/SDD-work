"""Build the 1-page, generic, company-data-free SDD pitch .docx.
No domain logic, no specs, no proprietary content. Needs: pip install python-docx
Run from anywhere: python3 scripts/build_pitch.py
"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(os.path.abspath(__file__))
REPO = os.path.dirname(HERE)
DIAG = os.path.join(REPO, "diagrams")

d = Document()

def h(text, size=14, color=(11, 37, 64), space_before=10):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(*color)
    return p

def body(text):
    p = d.add_paragraph(text); p.paragraph_format.space_after = Pt(6)
    for r in p.runs: r.font.size = Pt(10.5)
    return p

def pic(name, width):
    try:
        d.add_picture(os.path.join(DIAG, name), width=Inches(width))
        d.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        body(f"[{name} not embedded: {e}]")

title = d.add_paragraph()
r = title.add_run("Spec-Driven Development (SDD)"); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = RGBColor(11,37,64)
sub = d.add_paragraph(); r = sub.add_run("A brief: how we can run SDD in our own environment, on an internal AI platform")
r.italic = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(90,90,90)

h("The idea")
body("SDD makes the SPEC the single source of truth. A conceptual design becomes a precise "
     "spec (behaviour + acceptance criteria); the implementation is built to that spec with AI "
     "assistance; nothing is “done” until it is PROVEN against the spec. Everyone works from "
     "the spec, not from someone's code. The language or stack is irrelevant — the method is what matters.")

h("Why it is safe here — the discipline")
body("The core rule: access is not acceptance — acceptance is earned through proof. Every change "
     "is classified honestly (done / partially done / scaffolded / unproven) and must meet a proof "
     "bar appropriate to its type before it is called shipped. That is what makes AI-assisted "
     "delivery trustworthy in a regulated environment: speed without losing truth. The enemy we "
     "explicitly guard against is false completion — work claimed done but not proven.")

h("It fits our environment — no new infrastructure")
body("It runs on the internal, sanctioned AI platform we already have, inside the security boundary. "
     "No data leaves the building. No new model to procure. SDD is a method, not a tool stack — it "
     "rides on top of our existing versioning, CI/CD, containers, and vendor delivery, and composes "
     "with the BA → dev → QA pipeline rather than replacing any of it.")
pic("deploy.png", 4.3)

h("The method, one loop")
body("Orient (read the map) → write the spec with acceptance criteria → build to the spec → "
     "prove against the criteria → persist the decisions back into the docs so context compounds. Repeat.")
pic("loop.png", 2.6)

h("The MVP proposal (bounded)")
body("Take ONE well-bounded function for which a conceptual design already exists, derive its spec, "
     "and produce proven logic via SDD — small, self-contained, in days not months. A concrete "
     "demonstration that the method works inside our environment on a real, bounded piece of work. "
     "(The function and its spec come from the team — nothing assumed here.)")

h("The ask")
body("A small bounded slice + access to the internal AI platform for the demo. Outcome: a working, "
     "proven artifact plus a repeatable method the existing BA / dev / QA pipeline can adopt.")

out = os.path.join(REPO, "SDD_pitch.docx")
d.save(out)
print("WROTE", out)
